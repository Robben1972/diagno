from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import Hospital, Doctor, Chat, Message
from .serializers import HospitalSerializer, DoctorSerializer, ChatSerializer, MessageSerializer
from drf_spectacular.utils import extend_schema, OpenApiParameter
from math import radians, sin, cos, sqrt, atan2
from rest_framework.parsers import MultiPartParser, FormParser
from drf_spectacular.utils import OpenApiExample
import re

def calculate_distance(lat1, lon1, lat2, lon2):
    R = 6371  # Earth's radius in kilometers
    lat1, lon1, lat2, lon2 = map(radians, [lat1, lon1, lat2, lon2])
    dlat = lat2 - lat1
    dlon = lon2 - lon1
    a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlon/2)**2
    c = 2 * atan2(sqrt(a), sqrt(1-a))
    return R * c

class HospitalListView(APIView):
    @extend_schema(responses=HospitalSerializer(many=True))
    def get(self, request):
        hospitals = Hospital.objects.all()
        serializer = HospitalSerializer(hospitals, many=True)
        return Response(serializer.data)

    @extend_schema(request=HospitalSerializer, responses=HospitalSerializer)
    def post(self, request):
        serializer = HospitalSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class HospitalDetailView(APIView):
    @extend_schema(responses=HospitalSerializer)
    def get(self, request, pk):
        try:
            hospital = Hospital.objects.get(pk=pk)
        except Hospital.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        serializer = HospitalSerializer(hospital)
        return Response(serializer.data)

    @extend_schema(request=HospitalSerializer, responses=HospitalSerializer)
    def patch(self, request, pk):
        try:
            hospital = Hospital.objects.get(pk=pk)
        except Hospital.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        serializer = HospitalSerializer(hospital, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @extend_schema(responses={204: None})
    def delete(self, request, pk):
        try:
            hospital = Hospital.objects.get(pk=pk)
        except Hospital.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        hospital.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

class DoctorListView(APIView):
    @extend_schema(responses=DoctorSerializer(many=True))
    def get(self, request):
        doctors = Doctor.objects.all()
        serializer = DoctorSerializer(doctors, many=True)
        return Response(serializer.data)

    @extend_schema(request=DoctorSerializer, responses=DoctorSerializer)
    def post(self, request):
        serializer = DoctorSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class DoctorDetailView(APIView):
    @extend_schema(responses=DoctorSerializer)
    def get(self, request, pk):
        try:
            doctor = Doctor.objects.get(pk=pk)
        except Doctor.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        serializer = DoctorSerializer(doctor)
        return Response(serializer.data)

    @extend_schema(request=DoctorSerializer, responses=DoctorSerializer)
    def patch(self, request, pk):
        try:
            doctor = Doctor.objects.get(pk=pk)
        except Doctor.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        serializer = DoctorSerializer(doctor, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @extend_schema(responses={204: None})
    def delete(self, request, pk):
        try:
            doctor = Doctor.objects.get(pk=pk)
        except Doctor.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        doctor.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

class ChatListView(APIView):
    parser_classes = [MultiPartParser, FormParser]
    
    @extend_schema(
        responses=ChatSerializer(many=True),
        parameters=[
            OpenApiParameter(name='user_id', description='User ID to filter chats', required=True, type=str),
        ]
    )
    def get(self, request):
        user_id = request.query_params.get('user_id')
        if not user_id:
            return Response({"error": "user_id is required"}, status=status.HTTP_400_BAD_REQUEST)
        chats = Chat.objects.filter(user_id=user_id).prefetch_related('messages')
        serializer = ChatSerializer(chats, many=True)
        return Response(serializer.data)

    @extend_schema(
        request={
            'multipart/form-data': {
                'type': 'object',
                'properties': {
                    'user_id': {'type': 'string'},
                    'latitude': {'type': 'number'},
                    'longitude': {'type': 'number'},
                    'message': {'type': 'string', 'nullable': True},
                    'image': {'type': 'string', 'format': 'binary', 'nullable': True},
                    'file': {'type': 'string', 'format': 'binary', 'nullable': True},
                },
                'required': ['user_id', 'latitude', 'longitude']
            }
        },
        examples=[
            OpenApiExample(
                'Chat Example',
                value={
                    'user_id': '123',
                    'latitude': 41.3,
                    'longitude': 69.2,
                    'message': 'I have a rash',
                    'image': None,
                    'file': None
                },
                request_only=True,
                media_type='multipart/form-data'
            ),
        ],
        responses=ChatSerializer,
        description="Create a chat with user_id, latitude, longitude, and at least one of message, image, or file. AI will recommend doctors."
    )
    def post(self, request):
        serializer = ChatSerializer(data=request.data)
        if serializer.is_valid():
            user_id = serializer.validated_data.get('user_id')
            latitude = serializer.validated_data.get('latitude')
            longitude = serializer.validated_data.get('longitude')
            message = serializer.validated_data.get('message', '')
            image = serializer.validated_data.get('image')
            file = serializer.validated_data.get('file')

            if not (message or image or file):
                return Response(
                    {"error": "At least one of message, image, or file must be provided."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            image_file = image if image else None
            file_file = file if file else None

            doctors = Doctor.objects.all()
            doctor_info = "\n".join([
                f"{doc.id}. {doc.name} ({doc.field}) at {doc.hospital.name} [{doc.hospital.latitude}, {doc.hospital.longitude}]"
                for doc in doctors
            ])

            nearest_hospital = None
            min_distance = float('inf')
            for hospital in Hospital.objects.all():
                dist = calculate_distance(latitude, longitude, hospital.latitude, hospital.longitude)
                if dist < min_distance:
                    min_distance = dist
                    nearest_hospital = hospital

            prompt = (
                f"User message: {message}\n"
                f"Doctors available:\n{doctor_info}\n"
                f"Nearest hospital: {nearest_hospital.name if nearest_hospital else 'N/A'}\n"
                "Please recommend the most suitable doctor(s) for this user and answer the question. "
                "Include a list of recommended doctor IDs at the end of your response in the format: [1, 2, 3] or []"
            )

            file_text = None
            if file_file:
                name = getattr(file_file, 'name', str(file_file))
                if name.endswith('.txt'):
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        file_bytes = file_file.read()
                        file_text = "Text extracted from user's file: " + (file_bytes.decode('utf-8') if isinstance(file_bytes, bytes) else file_bytes)
                    else:
                        with open(file_file, 'r', encoding='utf-8') as f:
                            file_text = 'Text extracted from user\'s file: ' + f.read()
                elif name.endswith('.pdf'):
                    from PyPDF2 import PdfReader
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        reader = PdfReader(file_file)
                    else:
                        with open(file_file, 'rb') as f:
                            reader = PdfReader(f)
                    pdf_content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
                    file_text = pdf_content
                elif name.endswith('.docx'):
                    from docx import Document
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        doc = Document(file_file)
                    else:
                        doc = Document(file_file)
                    doc_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                    file_text = doc_content

            from .service.ai import generate_answer
            answer = generate_answer(prompt, image_file, file_text)
            if not answer:
                return Response({"error": "AI could not generate an answer."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Extract doctor IDs from the AI response
            doctor_ids = []
            match = re.search(r'\[([\d,\s]*)\]', answer)
            if match:
                doctor_ids_str = match.group(1)
                if doctor_ids_str:
                    doctor_ids = [int(id.strip()) for id in doctor_ids_str.split(',') if id.strip().isdigit()]

            chat = serializer.save()
            Message.objects.create(
                chat=chat,
                content=answer,
                is_from_user=False
            )
            return Response({"message": ''.join(answer.split('\n')[:-2]), "doctors": doctor_ids}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class ChatDetailView(APIView):
    parser_classes = [MultiPartParser, FormParser]

    @extend_schema(responses=ChatSerializer)
    def get(self, request, pk):
        try:
            chat = Chat.objects.get(pk=pk)
        except Chat.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        serializer = ChatSerializer(chat)
        return Response(serializer.data)

    @extend_schema(
        request={
            'multipart/form-data': {
                'type': 'object',
                'properties': {
                    'message': {'type': 'string', 'nullable': True},
                    'image': {'type': 'string', 'format': 'binary', 'nullable': True},
                    'file': {'type': 'string', 'format': 'binary', 'nullable': True},
                }
            }
        },
        examples=[
            OpenApiExample(
                'Chat Update Example',
                value={
                    'message': 'Here is an updated message',
                    'image': None,
                    'file': None
                },
                request_only=True,
                media_type='multipart/form-data'
            ),
        ],
        responses=ChatSerializer,
        description="Update a chat with message, image, or file. AI will use previous history and files/images to recommend doctors."
    )
    def patch(self, request, pk):
        try:
            chat = Chat.objects.get(pk=pk)
        except Chat.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        serializer = ChatSerializer(chat, data=request.data, partial=True)
        if serializer.is_valid():
            message = serializer.validated_data.get('message')
            image = serializer.validated_data.get('image')
            file = serializer.validated_data.get('file')

            if not (message or image or file):
                return Response(
                    {"error": "At least one of message, image, or file must be provided."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            image_file = image if image else None
            file_file = file if file else None

            doctors = Doctor.objects.all()
            doctor_info = "\n".join([
                f"{doc.id}. {doc.name} ({doc.field}) at {doc.hospital.name} [{doc.hospital.latitude}, {doc.hospital.longitude}]"
                for doc in doctors
            ])

            history = "\n".join([
                f"{'User' if msg.is_from_user else 'AI'}: {msg.content or '[file/image]'}"
                for msg in chat.messages.all()
            ])

            prompt = (
                f"Previous chat history:\n{history}\n"
                f"New user message: {message or '[file/image]'}\n"
                f"Doctors available:\n{doctor_info}\n"
                "Please recommend the most suitable doctor(s) for this user and answer the question. "
                "Include a list of recommended doctor IDs at the end of your response in the format: [1, 2, 3] or []"
            )

            file_text = None
            if file_file:
                name = getattr(file_file, 'name', str(file_file))
                if name.endswith('.txt'):
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        file_bytes = file_file.read()
                        file_text = "Text extracted from user's file: " + (file_bytes.decode('utf-8') if isinstance(file_bytes, bytes) else file_bytes)
                    else:
                        with open(file_file, 'r', encoding='utf-8') as f:
                            file_text = 'Text extracted from user\'s file: ' + f.read()
                elif name.endswith('.pdf'):
                    from PyPDF2 import PdfReader
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        reader = PdfReader(file_file)
                    else:
                        with open(file_file, 'rb') as f:
                            reader = PdfReader(f)
                    pdf_content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
                    file_text = pdf_content
                elif name.endswith('.docx'):
                    from docx import Document
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        doc = Document(file_file)
                    else:
                        doc = Document(file_file)
                    doc_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                    file_text = doc_content

            from .service.ai import generate_answer
            answer = generate_answer(prompt, image_file, file_text)
            if not answer:
                return Response({"error": "AI could not generate an answer."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Extract doctor IDs from the AI response
            doctor_ids = []
            match = re.search(r'\[([\d,\s]*)\]', answer)
            if match:
                doctor_ids_str = match.group(1)
                if doctor_ids_str:
                    doctor_ids = [int(id.strip()) for id in doctor_ids_str.split(',') if id.strip().isdigit()]

            Message.objects.create(
                chat=chat,
                content=message,
                image=image,
                file=file,
                is_from_user=True
            )
            Message.objects.create(
                chat=chat,
                content=answer,
                is_from_user=False
            )
            serializer.save()
            return Response({"message": ''.join(answer.split('\n')[:-2]), "doctors": doctor_ids})
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @extend_schema(responses={204: None})
    def delete(self, request, pk):
        try:
            chat = Chat.objects.get(pk=pk)
        except Chat.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        chat.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)