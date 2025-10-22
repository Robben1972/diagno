from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from doctors.models import Hospital, Doctor, Chat, Message
from .serializers import ChatSerializer
from drf_spectacular.utils import extend_schema, OpenApiParameter
from math import radians, sin, cos, sqrt, atan2
from rest_framework.parsers import MultiPartParser, FormParser
from drf_spectacular.utils import OpenApiExample
import re

def calculate_distance(lat1, lon1, lat2, lon2):
    R = 6371  # Earth's radius in kilometers
    lat1, lon1, lat2, lon2 = map(radians, [lat1, lon1, lat2, lon2])
    dlat = lat2 - lat1
    dlon = lon2 - lon1
    a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlon/2)**2
    c = 2 * atan2(sqrt(a), sqrt(1-a))
    return R * c

class ChatListView(APIView):
    parser_classes = [MultiPartParser, FormParser]

    @extend_schema(
        responses=ChatSerializer(many=True),
        parameters=[]  # Removed user_id parameter
    )
    def get(self, request):
        if not request.user.is_authenticated:
            return Response({"error": "Authentication required"}, status=status.HTTP_401_UNAUTHORIZED)
        user_id = request.user.id
        chats = Chat.objects.filter(user_id=user_id).prefetch_related('messages')
        serializer = ChatSerializer(chats, many=True)
        return Response(serializer.data)

    @extend_schema(
        request={
            'multipart/form-data': {
                'type': 'object',
                'properties': {
                    'latitude': {'type': 'number'},
                    'longitude': {'type': 'number'},
                    'message': {'type': 'string', 'nullable': True},
                    'image': {'type': 'string', 'format': 'binary', 'nullable': True},
                    'file': {'type': 'string', 'format': 'binary', 'nullable': True},
                },
                'required': ['latitude', 'longitude']
            }
        },
        examples=[
            OpenApiExample(
                'Chat Example',
                value={
                    'latitude': 41.3,
                    'longitude': 69.2,
                    'message': 'I have a rash',
                    'image': None,
                    'file': None
                },
                request_only=True,
                media_type='multipart/form-data'
            ),
        ],
        responses=ChatSerializer,
        description="Create a chat with latitude, longitude, and at least one of message, image, or file. AI will recommend doctors."
    )
    def post(self, request):
        if not request.user.is_authenticated:
            return Response({"error": "Authentication required"}, status=status.HTTP_401_UNAUTHORIZED)
        serializer = ChatSerializer(data=request.data, context={'request': request})
        if serializer.is_valid():
            latitude = serializer.validated_data.get('latitude')
            longitude = serializer.validated_data.get('longitude')
            message = serializer.validated_data.get('message', '')
            image = serializer.validated_data.get('image')
            file = serializer.validated_data.get('file')
            if not (message or image or file):
                return Response(
                    {"error": "At least one of message, image, or file must be provided."},
                    status=status.HTTP_400_BAD_REQUEST
                )
            image_file = image if image else None
            file_file = file if file else None
            doctors = Doctor.objects.all()
            doctor_info = "\n".join([
                f"{doc.id}. {doc.name} ({doc.field}) at {doc.hospital.name} [{doc.hospital.latitude}, {doc.hospital.longitude}]"
                for doc in doctors
            ])
            nearest_hospital = None
            min_distance = float('inf')
            for hospital in Hospital.objects.all():
                dist = calculate_distance(latitude, longitude, hospital.latitude, hospital.longitude)
                if dist < min_distance:
                    min_distance = dist
                    nearest_hospital = hospital
            prompt = (
                f"User message: {message}\n"
                f"Doctors available:\n{doctor_info}\n"
                f"Nearest hospital: {nearest_hospital.name if nearest_hospital else 'N/A'}\n"
            )
            file_text = None
            if file_file:
                name = getattr(file_file, 'name', str(file_file))
                if name.endswith('.txt'):
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        file_bytes = file_file.read()
                        file_text = "Text extracted from user's file: " + (file_bytes.decode('utf-8') if isinstance(file_bytes, bytes) else file_bytes)
                    else:
                        with open(file_file, 'r', encoding='utf-8') as f:
                            file_text = 'Text extracted from user\'s file: ' + f.read()
                elif name.endswith('.pdf'):
                    from PyPDF2 import PdfReader
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        reader = PdfReader(file_file)
                    else:
                        with open(file_file, 'rb') as f:
                            reader = PdfReader(f)
                    pdf_content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
                    file_text = pdf_content
                elif name.endswith('.docx'):
                    from docx import Document
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        doc = Document(file_file)
                    else:
                        doc = Document(file_file)
                    doc_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                    file_text = doc_content
            from doctors.service.ai import generate_answer
            response_text, doctor_ids = generate_answer(prompt, image_file, file_text)
            if not response_text:
                return Response({"error": "AI could not generate an answer."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            chat = serializer.save()

            Message.objects.create(
                chat=chat,
                content=response_text,
                is_from_user=False
            )

            return Response({
                "id": chat.id,
                "message": response_text,
                "doctors": doctor_ids
            }, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class ChatDetailView(APIView):
    parser_classes = [MultiPartParser, FormParser]

    @extend_schema(responses=ChatSerializer)
    def get(self, request, pk):
        try:
            chat = Chat.objects.get(pk=pk)
        except Chat.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        serializer = ChatSerializer(chat)
        return Response(serializer.data)

    @extend_schema(
        request={
            'multipart/form-data': {
                'type': 'object',
                'properties': {
                    'message': {'type': 'string', 'nullable': True},
                    'image': {'type': 'string', 'format': 'binary', 'nullable': True},
                    'file': {'type': 'string', 'format': 'binary', 'nullable': True},
                }
            }
        },
        examples=[
            OpenApiExample(
                'Chat Update Example',
                value={
                    'message': 'Here is an updated message',
                    'image': None,
                    'file': None
                },
                request_only=True,
                media_type='multipart/form-data'
            ),
        ],
        responses=ChatSerializer,
        description="Update a chat with message, image, or file. AI will use previous history and files/images to recommend doctors."
    )
    def patch(self, request, pk):
        try:
            chat = Chat.objects.get(pk=pk)
        except Chat.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        serializer = ChatSerializer(chat, data=request.data, partial=True)
        if serializer.is_valid():
            message = serializer.validated_data.get('message')
            image = serializer.validated_data.get('image')
            file = serializer.validated_data.get('file')
            if not (message or image or file):
                return Response(
                    {"error": "At least one of message, image, or file must be provided."},
                    status=status.HTTP_400_BAD_REQUEST
                )
            image_file = image if image else None
            file_file = file if file else None
            doctors = Doctor.objects.all()
            doctor_info = "\n".join([
                f"{doc.id}. {doc.name} ({doc.field}) at {doc.hospital.name} [{doc.hospital.latitude}, {doc.hospital.longitude}]"
                for doc in doctors
            ])
            history = "\n".join([
                f"{'User' if msg.is_from_user else 'AI'}: {msg.content or '[file/image]'}"
                for msg in chat.messages.all()
            ])
            prompt = (
                f"Previous chat history:\n{history}\n"
                f"New user message: {message or '[file/image]'}\n"
                f"Doctors available:\n{doctor_info}\n"
            )
            file_text = None
            if file_file:
                name = getattr(file_file, 'name', str(file_file))
                if name.endswith('.txt'):
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        file_bytes = file_file.read()
                        file_text = "Text extracted from user's file: " + (file_bytes.decode('utf-8') if isinstance(file_bytes, bytes) else file_bytes)
                    else:
                        with open(file_file, 'r', encoding='utf-8') as f:
                            file_text = 'Text extracted from user\'s file: ' + f.read()
                elif name.endswith('.pdf'):
                    from PyPDF2 import PdfReader
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        reader = PdfReader(file_file)
                    else:
                        with open(file_file, 'rb') as f:
                            reader = PdfReader(f)
                    pdf_content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
                    file_text = pdf_content
                elif name.endswith('.docx'):
                    from docx import Document
                    if hasattr(file_file, 'read'):
                        file_file.seek(0)
                        doc = Document(file_file)
                    else:
                        doc = Document(file_file)
                    doc_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                    file_text = doc_content
            from ...service.ai import generate_answer
            response_text, doctor_ids = generate_answer(prompt, image_file, file_text)
            if not response_text:
                return Response({"error": "AI could not generate an answer."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
           
            Message.objects.create(
                chat=chat,
                content=message,
                image=image,
                file=file,
                is_from_user=True
            )
            serializer.save()

            Message.objects.create(
                chat=chat,
                content=response_text,
                is_from_user=False
            )
            return Response({"id": chat.id, "message": ''.join(response_text), "doctors": doctor_ids})
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @extend_schema(responses={204: None})
    def delete(self, request, pk):
        try:
            chat = Chat.objects.get(pk=pk)
        except Chat.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        chat.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)