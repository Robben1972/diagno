import openai
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from PIL import Image
import os

OPENAI_API_KEY = "your-openai-api-key-here" 
openai.api_key = OPENAI_API_KEY

def generate_name(message: str) -> str:
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that generates concise titles."},
            {"role": "user", "content": f"Give the small title for this message in one line:\n {message}"}
        ],
        max_tokens=50
    )
    return response.choices[0].message.content.replace("\n", "").replace("\r", "").replace("*", "")

def generate_rag_prompt(query, context, history=""):
    escaped_context = context.replace("'", "").replace('"', "").replace("\n", " ")
    escaped_history = history.replace("'", "").replace('"', "").replace("\n", " ")
    prompt = ("""
            You are an intelligent, friendly, and helpful assistant. Your main job is to talk to a patient, answer their question, and give advice using the context provided.

            You should perform **TWO MAIN TASKS**:

            ---

            **TASK 1: Answer the patient's question clearly and kindly**

            - Use the CONTEXT and CHAT HISTORY to understand the situation.
            - Speak in simple, non-technical language that any patient can understand.
            - Avoid using medical jargon, programming terms, or developer-style explanations.
            - Explain important ideas in a friendly and supportive way.
            - If the CONTEXT is not useful, you can ignore it and answer based on your general knowledge.
            - Never say things like “You didn’t provide doctors” or “There is no doctor in this field.” Instead, just give helpful advice.

            ---

            **TASK 2: Suggest relevant medical specialists (if appropriate)**

            - If the question is about a medical condition, symptoms, or diagnosis, identify the types of medical specialists that might help the patient.
            - Use both the CONTEXT and the QUESTION to find clues about which specialists may be needed.
            - If no specialist is clearly relevant, just return an empty list.
            - Do **not** list general practitioners, unless it’s specifically requested.
            - If specialists are found, return them in **this exact format**:  
            `[Specialist1, Specialist2, ...]`  
            Example: `[Cardiologist, Neurologist]`

            - This list should appear at the **very end of the response**, and nothing else should come after it.

            ---

            **Additional Important Rules**:

            - You are talking to a patient — not a developer, not a doctor. Always be caring and conversational.
            - Avoid statements like: “I cannot help,” or “I don’t have data.”
            - Do not ask the patient to provide more information — just do your best with what is given.
            - Make sure your final output includes only one answer, with no duplicates or repeated messages.
            - Only list specialists if it is truly relevant. Otherwise, return `[]`.

            ---

            Now, use the following inputs to create your response:

            QUESTION: {query}

            CONTEXT: {escaped_context}

            CHAT HISTORY: {escaped_history}

            Now respond to the patient’s question. Be helpful, clear, and include the list of relevant specialists at the end (if any), formatted like: [Specialist1, Specialist2, ...]
            If no specialists are applicable, return an empty list like this: []
            """
    ).format(query=query, escaped_context=escaped_context, escaped_history=escaped_history)
    return prompt

def get_relevant_context_from_db(query):
    context = ""
    embedding_function = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    vector_db = Chroma(persist_directory="./healthy_documents", embedding_function=embedding_function)
    search_results = vector_db.similarity_search(query, k=6)
    for result in search_results:
        context += result.page_content + "\n"
    return context

def generate_answer(prompt, image_path=None, file_path=None):
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    content = [{"role": "user", "content": prompt}]
    
    if image_path:
        try:
            if hasattr(image_path, 'read'):
                img = Image.open(image_path)
            else:
                img = Image.open(str(image_path))
            # OpenAI API doesn't directly support PIL images; convert to base64 or skip for now
            print("Image processing not implemented for OpenAI API in this version.")
        except Exception as e:
            print(f"Error processing image: {e}")
    
    if file_path:
        try:
            name = getattr(file_path, 'name', str(file_path))
            if name.endswith('.txt'):
                if hasattr(file_path, 'read'):
                    file_path.seek(0)
                    content[0]["content"] += 'Text extracted from user\'s file: ' + (file_path.read().decode('utf-8') if hasattr(file_path, 'decode') else file_path.read())
                else:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content[0]["content"] += 'Text extracted from user\'s file: ' + f.read()
            elif name.endswith('.pdf'):
                from PyPDF2 import PdfReader
                if hasattr(file_path, 'read'):
                    file_path.seek(0)
                    reader = PdfReader(file_path)
                else:
                    with open(file_path, 'rb') as f:
                        reader = PdfReader(f)
                pdf_content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
                content.append({"role": "user", "content": pdf_content})
            elif name.endswith('.docx'):
                from docx import Document
                if hasattr(file_path, 'read'):
                    file_path.seek(0)
                    doc = Document(file_path)
                else:
                    doc = Document(file_path)
                doc_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                content.append({"role": "user", "content": doc_content})
            else:
                return None
        except Exception as e:
            print(f"Error processing file: {e}")
    
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful and informative bot."},
            *content
        ],
        max_tokens=1000
    )
    return response.choices[0].message.content

def get_answer(prompt: str, history: str, image_path=None, file_path=None):
    context = get_relevant_context_from_db(prompt)
    rag_prompt = generate_rag_prompt(query=prompt, context=context, history=history)
    answer = generate_answer(rag_prompt, image_path, file_path)
    return answer