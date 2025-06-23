import google.generativeai as genai
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
import os
import magic
from PIL import Image

GEMINI_API_KEY="AIzaSyBH-26V-1Jvsl2VwI6AcRwOtC3ZKWv1rD4"
genai.configure(api_key=GEMINI_API_KEY)

def generate_name(message: str) -> str:
    model = genai.GenerativeModel(model_name='gemini-2.0-flash')
    response = model.generate_content(f"Give the small title for this message in one line:\n {message}")
    return response.text.replace("\n", "").replace("\r", "").replace("*", "")

def generate_rag_prompt(query, context, history=""):
    escaped_context = context.replace("'", "").replace('"', "").replace("\n", " ")
    escaped_history = history.replace("'", "").replace('"', "").replace("\n", " ")
    prompt = ("""
        You are a helpful and informative bot that answers questions using text from the reference context included below. \
        Be sure to respond in a complete sentence, being comprehensive, including all relevant background information. \
        You are talking with a patient, so be sure to break down complicated concepts and \
        strike a friendly and conversational tone. \
        If the context is irrelevant to the answer, you may ignore it. If relevant, use the chat history to inform your response.
        
        Additionally, based on the question and context, identify the type of medical specialists who can help the patient. \
        Provide the list of specialists in the following format: [Specialist1, Specialist2, ...].
        
        QUESTION: '{query}'
        CONTEXT: '{escaped_context}'
        CHAT HISTORY: '{escaped_history}'
        
        ANSWER (include only the specialists list in the specified format if applicable):
        """).format(query=query, escaped_context=escaped_context, escaped_history=escaped_history)
    return prompt

def get_relevant_context_from_db(query):
    context = ""
    embedding_function = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    vector_db = Chroma(persist_directory="./healthy_documents", embedding_function=embedding_function)
    search_results = vector_db.similarity_search(query, k=6)
    for result in search_results:
        context += result.page_content + "\n"
    return context

def generate_answer(prompt, image_path=None, file_path=None):
    model = genai.GenerativeModel(model_name='gemini-2.0-flash')
    content = [prompt]
    
    if image_path:
        try:
            if hasattr(image_path, 'read'):
                img = Image.open(image_path)
            else:
                img = Image.open(str(image_path))
            content.append(img)
        except Exception as e:
            print(f"Error processing image: {e}")
    if file_path:
        try:
            name = getattr(file_path, 'name', str(file_path))
            if name.endswith('.txt'):
                if hasattr(file_path, 'read'):
                    file_path.seek(0)
                    content[0] += 'Text extracted from user\'s file: ' + (file_path.read().decode('utf-8') if hasattr(file_path, 'decode') else file_path.read())
                else:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content[0] += 'Text extracted from user\'s file: ' +  (f.read())
            elif name.endswith('.pdf'):
                from PyPDF2 import PdfReader
                if hasattr(file_path, 'read'):
                    file_path.seek(0)
                    reader = PdfReader(file_path)
                else:
                    with open(file_path, 'rb') as f:
                        reader = PdfReader(f)
                pdf_content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
                content.append(pdf_content)
            elif name.endswith('.docx'):
                from docx import Document
                if hasattr(file_path, 'read'):
                    file_path.seek(0)
                    doc = Document(file_path)
                else:
                    doc = Document(file_path)
                doc_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                content.append(doc_content)
            else:
                return None
        except Exception as e:
            print(f"Error processing file: {e}")
    
    answer = model.generate_content(content)
    return answer.text
def get_answer(prompt: str, history: str, image_path=None, file_path=None):
    context = get_relevant_context_from_db(prompt)
    rag_prompt = generate_rag_prompt(query=prompt, context=context, history=history)
    answer = generate_answer(rag_prompt, image_path, file_path)
    return answer